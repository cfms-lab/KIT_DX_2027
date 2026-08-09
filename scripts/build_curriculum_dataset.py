"""Build reusable CSV, SQLite, and Graphify corpus artifacts for DX curricula.

The extractor intentionally targets the 2025+ curriculum tables for the nine
DX departments in the official 2026 curriculum DOCX.  Source values and
normalized values are both retained so that corrections remain auditable.
"""

from __future__ import annotations

import csv
import re
import sqlite3
from collections import defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "src" / "★2026학년도 교육과정 운영기준 및 편성표(2026.02.27.).docx"
DATA_DIR = ROOT / "data"
CORPUS_DIR = ROOT / "graph_corpus"

SOURCE_TITLE = "2026학년도 교육과정 운영기준 및 편성표(2026.02.27.)"
COHORT_RULE = "2025학년도 이후 입학자 기준"


@dataclass(frozen=True)
class Department:
    department_id: str
    department_name: str
    is_engineering_accredited: int
    general_table_specs: tuple[tuple[int, str], ...]
    major_table_index: int


DEPARTMENTS = (
    Department("ARCH", "건축공학전공", 0, ((73, "dual_type_9"),), 74),
    Department("CIVIL", "토목공학전공", 0, ((113, "dual_type_9"),), 116),
    Department("IND", "산업공학전공", 0, ((228, "shared_type_8"),), 232),
    Department("MATHDATA", "수리빅데이터전공", 0, ((253, "shared_type_8"),), 256),
    Department("POLY", "고분자공학전공", 1, ((271, "shared_type_9_remark"),), 274),
    Department("MSE", "신소재공학전공", 1, ((291, "shared_type_8"), (292, "area_type_9")), 299),
    Department("MATDES", "소재디자인공학전공", 1, ((459, "dual_type_9"), (460, "dual_type_9")), 470),
    Department("CHEMENG", "화학공학전공", 1, ((492, "dual_type_9"), (493, "shared_type_8")), 501),
    Department("CHEMBIO", "화학생명소재전공", 0, ((517, "dual_type_9"), (518, "dual_type_9")), 521),
)


OFFERING_FIELDS = (
    "offering_id",
    "department_id",
    "department_name",
    "is_engineering_accredited",
    "curriculum_year",
    "cohort_rule",
    "curriculum_area",
    "is_msc",
    "year",
    "year_label",
    "semester",
    "requirement_type_raw",
    "requirement_type_normalized",
    "is_curriculum_required",
    "is_policy_required",
    "is_effectively_required",
    "required_reason",
    "course_code",
    "course_name_raw",
    "course_name_canonical",
    "is_design_course",
    "credit_structure_raw",
    "credits",
    "lecture_hours",
    "design_hours",
    "practice_hours",
    "source_document",
    "source_table_index",
    "source_row_index",
    "normalization_note",
    "quality_flag",
)


def compact(text: str) -> str:
    return " ".join(text.replace("\u00a0", " ").split())


def canonical_course_name(name: str, code: str) -> tuple[str, str]:
    raw_compact = compact(name)
    canonical = "".join(raw_compact.replace("*", "").split())
    changes: list[str] = []
    if any(character.isspace() for character in raw_compact):
        changes.append("공백 제거")
    if "*" in raw_compact:
        changes.append("원문 주석표시 * 제거")
    note = "; ".join(changes)
    if code == "LA0502" and canonical == "일반물리학실험":
        canonical = "일반물리학실험1"
        changes.append("LA0502 코드에 따라 일반물리학실험1로 정규화")
        note = "; ".join(changes)
    return canonical, note


def parse_credit(raw: str) -> tuple[int | None, int | None, int | None, int | None]:
    match = re.fullmatch(r"\s*(\d+)\s*-\s*(\d+)\s*-\s*(\d+)\s*-\s*(\d+)\s*", raw)
    if not match:
        return None, None, None, None
    return tuple(int(value) for value in match.groups())  # type: ignore[return-value]


def normalize_requirement(raw: str) -> tuple[str, int]:
    compacted = "".join(raw.split())
    if "필" in compacted:
        return "required", 1
    if "선" in compacted:
        return "elective", 0
    return "unspecified", 0


def parse_year(raw: str) -> tuple[int, str] | None:
    normalized = compact(raw)
    if normalized == "전학년":
        return 0, normalized
    match = re.match(r"([1-4])(?:학년)?$", normalized)
    if not match:
        return None
    return int(match.group(1)), normalized


def valid_course_code(code: str) -> bool:
    """Accept only actual curriculum course codes, not subtotal cell text."""
    return bool(re.fullmatch(r"[A-Z]{2}\d{4}", compact(code)))


def row_values(table, row_index: int) -> list[str]:
    return [compact(cell.text) for cell in table.rows[row_index].cells]


def general_slots(values: list[str], layout: str) -> list[tuple[int, str, str, str, str]]:
    """Return (semester, area_raw, requirement, code, name, credit) tuples."""
    if layout == "dual_type_9":
        return [
            (1, "", values[1], values[2], values[3], values[4]),
            (2, "", values[5], values[6], values[7], values[8]),
        ]
    if layout == "shared_type_8":
        return [
            (1, "", values[1], values[2], values[3], values[4]),
            (2, "", values[1], values[5], values[6], values[7]),
        ]
    if layout == "shared_type_9_remark":
        return [
            (1, "", values[1], values[2], values[3], values[4]),
            (2, "", values[1], values[5], values[6], values[7]),
        ]
    if layout == "area_type_9":
        return [
            (1, values[1], values[2], values[3], values[4], values[5]),
            (2, values[1], values[2], values[6], values[7], values[8]),
        ]
    raise ValueError(f"Unknown table layout: {layout}")


def classify_general_area(area_raw: str, requirement_raw: str) -> str:
    joined = "".join((area_raw + requirement_raw).split())
    if "학문기초" in joined or "지정" in joined:
        return "academic_foundation_msc"
    return "general_common"


def make_offering(
    *,
    sequence: int,
    department: Department,
    curriculum_area: str,
    year: int,
    year_label: str,
    semester: int,
    requirement_raw: str,
    course_code: str,
    course_name_raw: str,
    credit_raw: str,
    table_index: int,
    row_index: int,
    inferred_requirement: bool = False,
) -> dict[str, object]:
    canonical_name, normalization_note = canonical_course_name(course_name_raw, course_code)
    requirement_normalized, curriculum_required = normalize_requirement(requirement_raw)
    policy_required = int(
        bool(department.is_engineering_accredited)
        and canonical_name == "창의입문설계"
    )
    effective_required = int(bool(curriculum_required or policy_required))
    reasons: list[str] = []
    if curriculum_required:
        reasons.append("교육과정표 필수")
    if policy_required:
        reasons.append("공학교육인증 유지 전공의 창의입문설계 요건")
    credits, lecture, design, practice = parse_credit(credit_raw)
    flags: list[str] = []
    if inferred_requirement:
        flags.append("INFERRED_REQUIREMENT_FROM_PREVIOUS_ROW")
    if credits is None:
        flags.append("UNPARSED_CREDIT")
    if department.department_id == "ARCH" and course_code == "LA0508" and canonical_name == "일반물리학2":
        flags.append("POSSIBLE_SOURCE_CODE_TYPO")
    if normalization_note:
        flags.append("NORMALIZED_NAME")
    return {
        "offering_id": f"OFF{sequence:04d}",
        "department_id": department.department_id,
        "department_name": department.department_name,
        "is_engineering_accredited": department.is_engineering_accredited,
        "curriculum_year": 2026,
        "cohort_rule": COHORT_RULE,
        "curriculum_area": curriculum_area,
        "is_msc": int(curriculum_area == "academic_foundation_msc"),
        "year": year,
        "year_label": year_label,
        "semester": semester,
        "requirement_type_raw": requirement_raw,
        "requirement_type_normalized": requirement_normalized,
        "is_curriculum_required": curriculum_required,
        "is_policy_required": policy_required,
        "is_effectively_required": effective_required,
        "required_reason": "; ".join(reasons),
        "course_code": course_code,
        "course_name_raw": compact(course_name_raw),
        "course_name_canonical": canonical_name,
        "is_design_course": int("설계" in canonical_name),
        "credit_structure_raw": credit_raw,
        "credits": credits,
        "lecture_hours": lecture,
        "design_hours": design,
        "practice_hours": practice,
        "source_document": SOURCE_TITLE,
        "source_table_index": table_index,
        "source_row_index": row_index,
        "normalization_note": normalization_note,
        "quality_flag": ";".join(flags),
    }


def extract_offerings() -> list[dict[str, object]]:
    doc = Document(SOURCE)
    offerings: list[dict[str, object]] = []
    sequence = 1
    for department in DEPARTMENTS:
        for table_index, layout in department.general_table_specs:
            table = doc.tables[table_index]
            for row_index in range(2, len(table.rows)):
                values = row_values(table, row_index)
                parsed_year = parse_year(values[0])
                if parsed_year is None:
                    continue
                year, year_label = parsed_year
                for semester, area_raw, requirement, code, name, credit in general_slots(values, layout):
                    if not valid_course_code(code) or not name:
                        continue
                    curriculum_area = classify_general_area(area_raw, requirement)
                    offerings.append(
                        make_offering(
                            sequence=sequence,
                            department=department,
                            curriculum_area=curriculum_area,
                            year=year,
                            year_label=year_label,
                            semester=semester,
                            requirement_raw=requirement,
                            course_code=code,
                            course_name_raw=name,
                            credit_raw=credit,
                            table_index=table_index,
                            row_index=row_index,
                        )
                    )
                    sequence += 1

        table_index = department.major_table_index
        table = doc.tables[table_index]
        last_requirement = ""
        for row_index in range(2, len(table.rows)):
            values = row_values(table, row_index)
            parsed_year = parse_year(values[0])
            if parsed_year is None:
                continue
            year, year_label = parsed_year
            raw_requirement = values[1]
            inferred = False
            if re.search(r"전\s*[필선]", raw_requirement):
                last_requirement = raw_requirement
            elif any(values[col] for col in (2, 3, 5, 6)) and last_requirement:
                raw_requirement = last_requirement
                inferred = True
            else:
                continue
            for semester, code_col, name_col, credit_col in ((1, 2, 3, 4), (2, 5, 6, 7)):
                code, name, credit = values[code_col], values[name_col], values[credit_col]
                if not valid_course_code(code) or not name:
                    continue
                offerings.append(
                    make_offering(
                        sequence=sequence,
                        department=department,
                        curriculum_area="major",
                        year=year,
                        year_label=year_label,
                        semester=semester,
                        requirement_raw=raw_requirement,
                        course_code=code,
                        course_name_raw=name,
                        credit_raw=credit,
                        table_index=table_index,
                        row_index=row_index,
                        inferred_requirement=inferred,
                    )
                )
                sequence += 1
    return offerings


def write_csv(path: Path, rows: list[dict[str, object]], fieldnames: tuple[str, ...] | list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def build_overlap(offerings: list[dict[str, object]]) -> list[dict[str, object]]:
    buckets: dict[str, list[dict[str, object]]] = defaultdict(list)
    for offering in offerings:
        if int(offering["year"]) == 0:
            continue
        buckets[str(offering["course_name_canonical"])].append(offering)
    rows: list[dict[str, object]] = []
    for course_name, items in sorted(buckets.items()):
        departments_any = sorted({str(item["department_name"]) for item in items})
        curriculum_required = sorted({
            str(item["department_name"]) for item in items if int(item["is_curriculum_required"])
        })
        effective_required = sorted({
            str(item["department_name"]) for item in items if int(item["is_effectively_required"])
        })
        terms = sorted({
            f"{item['department_name']}:{item['year']}-{item['semester']}" for item in items
        })
        rows.append({
            "course_name_canonical": course_name,
            "department_count_any": len(departments_any),
            "department_count_curriculum_required": len(curriculum_required),
            "department_count_effective_required": len(effective_required),
            "departments_any": "; ".join(departments_any),
            "departments_curriculum_required": "; ".join(curriculum_required),
            "departments_effective_required": "; ".join(effective_required),
            "department_terms": "; ".join(terms),
        })
    return rows


def build_quality_issues(offerings: list[dict[str, object]]) -> list[dict[str, object]]:
    issues: list[dict[str, object]] = []
    issue_no = 1
    for offering in offerings:
        flags = str(offering["quality_flag"])
        if "POSSIBLE_SOURCE_CODE_TYPO" in flags:
            issues.append({
                "issue_id": f"Q{issue_no:03d}",
                "severity": "warning",
                "issue_type": "possible_source_code_typo",
                "department_name": offering["department_name"],
                "course_code": offering["course_code"],
                "course_name": offering["course_name_raw"],
                "detail": "건축공학 표의 LA0508 일반물리학2는 타 전공의 LA0506 일반물리학2 및 LA0508 일반화학2와 충돌함. 원문 확인 필요.",
                "source_table_index": offering["source_table_index"],
                "source_row_index": offering["source_row_index"],
            })
            issue_no += 1
        if "LA0502 코드" in str(offering["normalization_note"]):
            issues.append({
                "issue_id": f"Q{issue_no:03d}",
                "severity": "info",
                "issue_type": "normalized_course_name",
                "department_name": offering["department_name"],
                "course_code": offering["course_code"],
                "course_name": offering["course_name_raw"],
                "detail": offering["normalization_note"],
                "source_table_index": offering["source_table_index"],
                "source_row_index": offering["source_row_index"],
            })
            issue_no += 1

    code_names: dict[str, set[str]] = defaultdict(set)
    name_codes: dict[str, set[str]] = defaultdict(set)
    for offering in offerings:
        code_names[str(offering["course_code"])].add(str(offering["course_name_canonical"]))
        name_codes[str(offering["course_name_canonical"])].add(str(offering["course_code"]))
    for code, names in sorted(code_names.items()):
        if len(names) <= 1 or code == "FP0001":
            continue
        issues.append({
            "issue_id": f"Q{issue_no:03d}",
            "severity": "warning",
            "issue_type": "course_code_multiple_names",
            "department_name": "",
            "course_code": code,
            "course_name": "; ".join(sorted(names)),
            "detail": "동일 과목코드가 둘 이상의 정규화 과목명으로 나타남.",
            "source_table_index": "",
            "source_row_index": "",
        })
        issue_no += 1
    for name, codes in sorted(name_codes.items()):
        if len(codes) <= 1:
            continue
        issues.append({
            "issue_id": f"Q{issue_no:03d}",
            "severity": "info",
            "issue_type": "course_name_multiple_codes",
            "department_name": "",
            "course_code": "; ".join(sorted(codes)),
            "course_name": name,
            "detail": "동일 정규화 과목명이 둘 이상의 코드로 개설됨. 동일인정 여부는 별도 규정 확인 필요.",
            "source_table_index": "",
            "source_row_index": "",
        })
        issue_no += 1
    return issues


def build_database(
    offerings: list[dict[str, object]],
    overlaps: list[dict[str, object]],
    quality_issues: list[dict[str, object]],
) -> None:
    output = DATA_DIR / "dx_curriculum.sqlite"
    temporary = DATA_DIR / "dx_curriculum.sqlite.tmp"
    temporary.unlink(missing_ok=True)
    connection = sqlite3.connect(temporary)
    connection.executescript(
        """
        PRAGMA foreign_keys = ON;
        CREATE TABLE departments (
            department_id TEXT PRIMARY KEY,
            department_name TEXT NOT NULL UNIQUE,
            is_engineering_accredited INTEGER NOT NULL CHECK (is_engineering_accredited IN (0,1))
        );
        CREATE TABLE course_offerings (
            offering_id TEXT PRIMARY KEY,
            department_id TEXT NOT NULL REFERENCES departments(department_id),
            department_name TEXT NOT NULL,
            is_engineering_accredited INTEGER NOT NULL,
            curriculum_year INTEGER NOT NULL,
            cohort_rule TEXT NOT NULL,
            curriculum_area TEXT NOT NULL,
            is_msc INTEGER NOT NULL,
            year INTEGER NOT NULL,
            year_label TEXT NOT NULL,
            semester INTEGER NOT NULL CHECK (semester IN (1,2)),
            requirement_type_raw TEXT NOT NULL,
            requirement_type_normalized TEXT NOT NULL,
            is_curriculum_required INTEGER NOT NULL,
            is_policy_required INTEGER NOT NULL,
            is_effectively_required INTEGER NOT NULL,
            required_reason TEXT NOT NULL,
            course_code TEXT NOT NULL,
            course_name_raw TEXT NOT NULL,
            course_name_canonical TEXT NOT NULL,
            is_design_course INTEGER NOT NULL,
            credit_structure_raw TEXT NOT NULL,
            credits INTEGER,
            lecture_hours INTEGER,
            design_hours INTEGER,
            practice_hours INTEGER,
            source_document TEXT NOT NULL,
            source_table_index INTEGER NOT NULL,
            source_row_index INTEGER NOT NULL,
            normalization_note TEXT NOT NULL,
            quality_flag TEXT NOT NULL
        );
        CREATE TABLE course_overlap (
            course_name_canonical TEXT PRIMARY KEY,
            department_count_any INTEGER NOT NULL,
            department_count_curriculum_required INTEGER NOT NULL,
            department_count_effective_required INTEGER NOT NULL,
            departments_any TEXT NOT NULL,
            departments_curriculum_required TEXT NOT NULL,
            departments_effective_required TEXT NOT NULL,
            department_terms TEXT NOT NULL
        );
        CREATE TABLE quality_issues (
            issue_id TEXT PRIMARY KEY,
            severity TEXT NOT NULL,
            issue_type TEXT NOT NULL,
            department_name TEXT NOT NULL,
            course_code TEXT NOT NULL,
            course_name TEXT NOT NULL,
            detail TEXT NOT NULL,
            source_table_index TEXT NOT NULL,
            source_row_index TEXT NOT NULL
        );
        CREATE INDEX idx_offerings_dept_year_term ON course_offerings(department_id, year, semester);
        CREATE INDEX idx_offerings_course ON course_offerings(course_name_canonical);
        CREATE INDEX idx_offerings_msc_required ON course_offerings(is_msc, is_effectively_required, year);
        CREATE VIEW v_first_year_effective_required AS
            SELECT department_name, semester, course_code, course_name_canonical, credits,
                   curriculum_area, required_reason
            FROM course_offerings
            WHERE year = 1 AND is_effectively_required = 1;
        CREATE VIEW v_required_course_overlap AS
            SELECT course_name_canonical,
                   COUNT(DISTINCT department_id) AS department_count,
                   GROUP_CONCAT(DISTINCT department_name) AS departments
            FROM course_offerings
            WHERE is_effectively_required = 1 AND year BETWEEN 1 AND 4
            GROUP BY course_name_canonical
            HAVING COUNT(DISTINCT department_id) >= 2;
        CREATE VIEW v_later_msc_required_overlap AS
            SELECT course_name_canonical,
                   COUNT(DISTINCT department_id) AS department_count,
                   GROUP_CONCAT(DISTINCT department_name) AS departments,
                   GROUP_CONCAT(DISTINCT year || '-' || semester) AS year_terms
            FROM course_offerings
            WHERE is_msc = 1 AND is_effectively_required = 1 AND year BETWEEN 2 AND 4
            GROUP BY course_name_canonical
            HAVING COUNT(DISTINCT department_id) >= 2;
        """
    )
    connection.executemany(
        "INSERT INTO departments VALUES (?,?,?)",
        [(d.department_id, d.department_name, d.is_engineering_accredited) for d in DEPARTMENTS],
    )
    placeholders = ",".join("?" for _ in OFFERING_FIELDS)
    connection.executemany(
        f"INSERT INTO course_offerings ({','.join(OFFERING_FIELDS)}) VALUES ({placeholders})",
        [[row[field] for field in OFFERING_FIELDS] for row in offerings],
    )
    overlap_fields = list(overlaps[0].keys())
    connection.executemany(
        f"INSERT INTO course_overlap ({','.join(overlap_fields)}) VALUES ({','.join('?' for _ in overlap_fields)})",
        [[row[field] for field in overlap_fields] for row in overlaps],
    )
    issue_fields = list(quality_issues[0].keys()) if quality_issues else []
    if issue_fields:
        connection.executemany(
            f"INSERT INTO quality_issues ({','.join(issue_fields)}) VALUES ({','.join('?' for _ in issue_fields)})",
            [[row[field] for field in issue_fields] for row in quality_issues],
        )
    connection.commit()
    integrity = connection.execute("PRAGMA integrity_check").fetchone()[0]
    if integrity != "ok":
        raise RuntimeError(f"SQLite integrity check failed: {integrity}")
    connection.close()
    temporary.replace(output)


def build_graph_corpus(offerings: list[dict[str, object]], overlaps: list[dict[str, object]]) -> None:
    CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    department_dir = CORPUS_DIR / "departments"
    department_dir.mkdir(parents=True, exist_ok=True)
    summary_lines = [
        "# DX대학 교육과정 지식 그래프 코퍼스",
        "",
        f"- 기준 문서: {SOURCE_TITLE}",
        f"- 적용 기준: {COHORT_RULE}",
        "- 대상 전공: 9개",
        "- 관계 판정: 원문 과목명과 코드, 정규화 과목명을 함께 보존",
        "",
        "## 공통 필수 교과목",
        "",
    ]
    for row in overlaps:
        if int(row["department_count_effective_required"]) >= 2:
            summary_lines.append(
                f"- {row['course_name_canonical']}: 유효 필수 {row['department_count_effective_required']}개 전공 "
                f"({row['departments_effective_required']})"
            )
    (CORPUS_DIR / "overview.md").write_text("\n".join(summary_lines) + "\n", encoding="utf-8")

    by_department: dict[str, list[dict[str, object]]] = defaultdict(list)
    for offering in offerings:
        by_department[str(offering["department_id"])].append(offering)
    for department in DEPARTMENTS:
        lines = [
            f"# {department.department_name}",
            "",
            f"- 전공 ID: {department.department_id}",
            f"- 공학교육인증 유지: {'예' if department.is_engineering_accredited else '아니오'}",
            f"- 적용 기준: {COHORT_RULE}",
            "",
            "## 교과목",
            "",
            "| 학년-학기 | 영역 | 이수구분 | 코드 | 정규화 과목명 | 원문 과목명 | 학점 | 필수 근거 |",
            "|---|---|---|---|---|---|---:|---|",
        ]
        items = sorted(by_department[department.department_id], key=lambda r: (int(r["year"]), int(r["semester"]), str(r["course_code"])))
        for item in items:
            year_term = f"{item['year_label']}-{item['semester']}"
            lines.append(
                f"| {year_term} | {item['curriculum_area']} | {item['requirement_type_raw']} | "
                f"{item['course_code']} | {item['course_name_canonical']} | {item['course_name_raw']} | "
                f"{item['credits'] if item['credits'] is not None else ''} | {item['required_reason']} |"
            )
        filename = f"{department.department_id}_{department.department_name}.md"
        (department_dir / filename).write_text("\n".join(lines) + "\n", encoding="utf-8")

    regulation_source = ROOT / "src" / "국립금오공과대학교_학사운영_규정_전문202606.docx"
    if regulation_source.exists():
        regulation_dir = CORPUS_DIR / "regulations"
        regulation_dir.mkdir(parents=True, exist_ok=True)
        regulation_doc = Document(regulation_source)
        regulation_lines = [
            "# 국립금오공과대학교 학사운영 규정",
            "",
            f"- 원본 파일: {regulation_source.name}",
            "- 그래프 탐색용 변환본이며 공식 원문을 대체하지 않음",
            "",
        ]
        for child in regulation_doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                text = compact(Paragraph(child, regulation_doc).text)
                if not text:
                    continue
                if re.match(r"제\d+장", text):
                    regulation_lines.extend([f"## {text}", ""])
                elif re.match(r"제\d+절", text):
                    regulation_lines.extend([f"### {text}", ""])
                elif re.match(r"제\d+조", text):
                    regulation_lines.extend([f"#### {text}", ""])
                else:
                    regulation_lines.extend([text, ""])
            elif isinstance(child, CT_Tbl):
                table = Table(child, regulation_doc)
                for row in table.rows:
                    cells = [compact(cell.text).replace("|", "\\|") for cell in row.cells]
                    # Most regulation tables are blank forms with heavily merged cells.
                    # Keep only rows relevant to curriculum/major/credit rules and
                    # collapse repeated merged-cell text before graph extraction.
                    joined = " ".join(cells)
                    if not any(keyword in joined for keyword in ("교육과정", "전공", "학점", "졸업", "이수", "자율전공")):
                        continue
                    unique_cells: list[str] = []
                    for cell_text in cells:
                        if cell_text and (not unique_cells or unique_cells[-1] != cell_text):
                            unique_cells.append(cell_text)
                    regulation_lines.append("- 표: " + " / ".join(unique_cells))
                regulation_lines.append("")
        (regulation_dir / "학사운영_규정_202606.md").write_text(
            "\n".join(regulation_lines) + "\n", encoding="utf-8"
        )


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    offerings = extract_offerings()
    overlaps = build_overlap(offerings)
    quality_issues = build_quality_issues(offerings)

    write_csv(
        DATA_DIR / "departments.csv",
        [
            {
                "department_id": d.department_id,
                "department_name": d.department_name,
                "is_engineering_accredited": d.is_engineering_accredited,
            }
            for d in DEPARTMENTS
        ],
        ["department_id", "department_name", "is_engineering_accredited"],
    )
    write_csv(DATA_DIR / "course_offerings.csv", offerings, list(OFFERING_FIELDS))
    write_csv(DATA_DIR / "course_overlap.csv", overlaps, list(overlaps[0].keys()))
    if quality_issues:
        write_csv(DATA_DIR / "quality_issues.csv", quality_issues, list(quality_issues[0].keys()))
    build_database(offerings, overlaps, quality_issues)
    build_graph_corpus(offerings, overlaps)

    print(f"departments={len(DEPARTMENTS)}")
    print(f"offerings={len(offerings)}")
    print(f"distinct_courses={len(overlaps)}")
    print(f"quality_issues={len(quality_issues)}")


if __name__ == "__main__":
    main()
