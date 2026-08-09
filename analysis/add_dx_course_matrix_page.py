"""Append a landscape A4 course-by-department matrix as page 2.

The existing one-page brief is opened as the source and is never overwritten.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "output" / "DX대학_자율전공_1학년_교육과정_회의브리프_1p.docx"
OUTPUT = ROOT / "output" / "DX대학_자율전공_1학년_교육과정_회의브리프_2p.docx"

NAVY = "17365D"
BLUE = "2E74B5"
TEXT_GRAY = "4A5560"
GRID = "CAD1D9"
WHITE = "FFFFFF"

GROUPS = [
    ("9개 전공 공통", "D9EAD3", "38761D", [
        ("글쓰기와\n발표", "글쓰기와발표"),
        ("글로벌\n커뮤니케이션", "글로벌커뮤니케이션"),
        ("디지털\n문해력", "디지털문해력"),
        ("대학\n수학1", "대학수학1"),
        ("대학\n수학2", "대학수학2"),
    ]),
    ("기초과학", "DCEAF7", "1F4E78", [
        ("일반\n물리학1", "일반물리학1"),
        ("일반\n물리학2", "일반물리학2"),
        ("일반\n화학1", "일반화학1"),
        ("일반\n화학2", "일반화학2"),
        ("물리학\n실험1", "일반물리학실험1"),
        ("물리학\n실험2", "일반물리학실험2"),
        ("화학\n실험1", "일반화학실험1"),
        ("화학\n실험2", "일반화학실험2"),
    ]),
    ("설계·데이터", "FCE5CD", "9C5700", [
        ("창의입문\n설계", "창의입문설계"),
        ("컴퓨터\n프로그래밍", "컴퓨터프로그래밍언어"),
        ("확률 및\n통계", "확률및통계"),
    ]),
    ("전공·계열 특화", "EDEDED", "4A5560", [
        ("건축토목환경\n공학개론", "건축토목환경공학개론"),
        ("스마트그린\n빌딩의이해", "스마트그린빌딩의이해"),
        ("건축과\n컴퓨터", "건축과컴퓨터"),
        ("공업\n역학", "공업역학"),
        ("스마트제조\n개론", "스마트제조개론"),
        ("빅데이터의\n세계", "빅데이터의세계"),
    ]),
]


DEPARTMENTS = [
    ("건축공학전공", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 1, "디지털문해력": 2,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "건축토목환경공학개론": 1,
        "스마트그린빌딩의이해": 2, "건축과컴퓨터": 2,
    }),
    ("토목공학전공", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 1, "디지털문해력": 2,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "컴퓨터프로그래밍언어": 2,
        "건축토목환경공학개론": 1, "공업역학": 2,
    }),
    ("산업공학전공", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 2,
        "대학수학1": 1, "대학수학2": 2, "컴퓨터프로그래밍언어": 1,
        "확률및통계": 1, "스마트제조개론": 2, "빅데이터의세계": 2,
    }),
    ("수리빅데이터전공", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 2,
        "대학수학1": 1, "대학수학2": 2, "컴퓨터프로그래밍언어": 1,
        "확률및통계": 1, "빅데이터의세계": 2,
    }),
    ("고분자공학전공 (인증)", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 1,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "일반화학2": 2,
        "일반물리학실험1": 1, "일반물리학실험2": 2,
        "일반화학실험1": 1, "일반화학실험2": 2, "창의입문설계": 2,
    }),
    ("신소재공학전공 (인증)", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 1,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "일반화학2": 2,
        "일반물리학실험1": 1, "일반물리학실험2": 2,
        "일반화학실험1": 1, "일반화학실험2": 2, "창의입문설계": 2,
    }),
    ("소재디자인공학전공 (인증)", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 1,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "일반화학2": 2,
        "일반물리학실험1": 1, "일반물리학실험2": 2,
        "일반화학실험1": 1, "일반화학실험2": 2, "창의입문설계": 2,
    }),
    ("화학공학전공 (인증)", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 1,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "일반화학2": 2,
        "일반물리학실험1": 1, "일반물리학실험2": 2,
        "일반화학실험1": 1, "일반화학실험2": 2, "창의입문설계": 2,
    }),
    ("화학생명소재전공", {
        "글쓰기와발표": 1, "글로벌커뮤니케이션": 2, "디지털문해력": 1,
        "대학수학1": 1, "대학수학2": 2, "일반물리학1": 1,
        "일반물리학2": 2, "일반화학1": 1, "일반화학2": 2,
        "일반물리학실험1": 1, "일반물리학실험2": 2,
        "일반화학실험1": 1, "일반화학실험2": 2, "창의입문설계": 2,
    }),
]


def set_font(run, size=8, bold=False, color="000000"):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=50, start=50, bottom=50, end=50):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=GRID, size=5):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def geometry(table, widths, indent=0):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(cell)


def write(cell, text, *, size=7.1, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.92
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def unlink_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def build():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document(SOURCE)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(9)
    section.bottom_margin = Mm(9)
    section.left_margin = Mm(9)
    section.right_margin = Mm(9)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)
    unlink_header_footer(section)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    write_run = header.add_run("DX대학 자율전공 교육과정 | 교과목 교집합표")
    set_font(write_run, size=8, bold=True, color=TEXT_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    write_run = footer.add_run("근거: 2026학년도 교육과정 운영기준 및 편성표  |  2026.08.09  |  2 / 2")
    set_font(write_run, size=7.5, color=TEXT_GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("DX대학 9개 전공의 1학년 교과목 교집합")
    set_font(run, size=15.5, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("행은 전공, 열은 교과목입니다. 숫자는 이수 학기이며, 빈칸은 해당 전공의 1학년 고정 교육과정에 없음을 뜻합니다.")
    set_font(run, size=8.5, color=TEXT_GRAY)

    course_defs = [course for _, _, _, courses in GROUPS for course in courses]
    ncols = 1 + len(course_defs)
    table = doc.add_table(rows=2 + len(DEPARTMENTS) + 1, cols=ncols)
    dept_width = 1900
    course_width = 625
    widths = [dept_width] + [course_width] * len(course_defs)
    geometry(table, widths)
    borders(table)

    # Two-row header, with group bands on top.
    dept_header = table.cell(0, 0).merge(table.cell(1, 0))
    shade(dept_header, NAVY)
    write(dept_header, "학과(전공)", size=8, bold=True, color=WHITE)

    col_index = 1
    course_meta = []
    for group_name, fill, color, courses in GROUPS:
        start = col_index
        end = col_index + len(courses) - 1
        group_cell = table.cell(0, start)
        if end > start:
            group_cell = group_cell.merge(table.cell(0, end))
        shade(group_cell, fill)
        write(group_cell, group_name, size=7.4, bold=True, color=color)
        for short_name, key in courses:
            cell = table.cell(1, col_index)
            shade(cell, fill)
            write(cell, short_name, size=6.6, bold=True, color=color)
            course_meta.append((key, fill, color))
            col_index += 1

    # Department rows.
    for row_index, (department, subjects) in enumerate(DEPARTMENTS, start=2):
        dept_cell = table.cell(row_index, 0)
        shade(dept_cell, "F4F6F8" if row_index % 2 == 0 else WHITE)
        write(dept_cell, department, size=7.3, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT)
        for offset, (key, fill, color) in enumerate(course_meta, start=1):
            cell = table.cell(row_index, offset)
            semester = subjects.get(key)
            if semester is not None:
                shade(cell, fill)
                write(cell, str(semester), size=7.6, bold=True, color=color)
            else:
                shade(cell, "FAFBFC" if row_index % 2 == 0 else WHITE)
                write(cell, "", size=7)

    # Frequency row makes the intersection immediately explicit.
    freq_row = 2 + len(DEPARTMENTS)
    shade(table.cell(freq_row, 0), NAVY)
    write(table.cell(freq_row, 0), "포함 전공 수", size=7.5, bold=True, color=WHITE)
    for offset, (key, fill, color) in enumerate(course_meta, start=1):
        count = sum(1 for _, subjects in DEPARTMENTS if key in subjects)
        shade(table.cell(freq_row, offset), fill)
        write(table.cell(freq_row, offset), f"{count}/9", size=7.1, bold=True, color=color)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("읽는 법  ")
    set_font(run, size=8, bold=True, color=BLUE)
    run = note.add_run("초록색 5개 열은 9개 전공의 완전 교집합(글쓰기와발표·글로벌커뮤니케이션·디지털문해력·대학수학1·2)입니다. ")
    set_font(run, size=8, color=TEXT_GRAY)
    run = note.add_run("(인증)")
    set_font(run, size=8, bold=True, color="9C5700")
    run = note.add_run("은 공학인증 유지 전공이며, 창의입문설계는 화학생명소재전공까지 포함해 5/9 전공에 편성되어 있습니다.")
    set_font(run, size=8, color=TEXT_GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
