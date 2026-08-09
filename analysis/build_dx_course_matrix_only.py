"""Rebuild the second deliverable as a standalone one-page A4 matrix."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

import add_dx_course_matrix_page as base


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "DX대학_자율전공_1학년_교육과정_회의브리프_2p.docx"

NAVY = base.NAVY
BLUE = base.BLUE
TEXT_GRAY = base.TEXT_GRAY
WHITE = base.WHITE
GRID = base.GRID

COMMON_FILL = "D9EAD3"
COMMON_DARK = "38761D"
SCIENCE_FILL = "DCEAF7"
SCIENCE_DARK = "1F4E78"
DESIGN_FILL = "FCE5CD"
DESIGN_DARK = "9C5700"
MSC_FILL = "FFF2CC"
MSC_DARK = "7F6000"
SPECIAL_FILL = "EDEDED"
SPECIAL_DARK = "4A5560"
ACCREDITED_BORDER = DESIGN_DARK  # Match the 창의입문설계 font color.


COMMON = [
    ("글쓰기와\n발표", "글쓰기와발표"),
    ("글로벌\n커뮤니케이션", "글로벌커뮤니케이션"),
    ("디지털\n문해력", "디지털문해력"),
    ("대학\n수학1", "대학수학1"),
    ("대학\n수학2", "대학수학2"),
]

SCIENCE = [
    ("일반\n물리학1", "일반물리학1"),
    ("일반\n물리학2", "일반물리학2"),
    ("일반\n화학1", "일반화학1"),
    ("일반\n화학2", "일반화학2"),
    ("물리학\n실험1", "일반물리학실험1"),
    ("물리학\n실험2", "일반물리학실험2"),
    ("화학\n실험1", "일반화학실험1"),
    ("화학\n실험2", "일반화학실험2"),
]

DESIGN = [("창의입문\n설계", "창의입문설계")]
MSC = [
    ("컴퓨터\n프로그래밍", "컴퓨터프로그래밍언어"),
    ("확률 및\n통계", "확률및통계"),
    ("고급\n프로그래밍\n언어", "고급프로그래밍언어"),
    ("공학\n수학", "공학수학"),
    ("공학\n수학1", "공학수학1"),
    ("공학\n수학2", "공학수학2"),
]

SPECIAL_KEYS = [
    "건축토목환경공학개론",
    "스마트그린빌딩의이해",
    "건축과컴퓨터",
    "공업역학",
    "스마트제조개론",
    "빅데이터의세계",
]

# Exact-name matches found as required courses after the first year in the
# 2025+ department-specific curriculum tables. Values are year-semester.
LATER_REQUIRED = {
    "건축공학전공": {"확률및통계": "2-2", "공학수학": "2-1"},
    "토목공학전공": {"확률및통계": "2-2", "공학수학": "2-1"},
    "산업공학전공": {"고급프로그래밍언어": "2-1", "공학수학": "2-2"},
    "수리빅데이터전공": {"고급프로그래밍언어": "2-2"},
    "고분자공학전공 (인증)": {"공학수학1": "2-1", "공학수학2": "2-2"},
    "신소재공학전공 (인증)": {"공학수학1": "2-1", "공학수학2": "2-2"},
    "소재디자인공학전공 (인증)": {"확률및통계": "2-1", "고급프로그래밍언어": "2-2"},
    "화학공학전공 (인증)": {"공학수학1": "2-1", "공학수학2": "2-2"},
}

# Present in year 1 but not required. The four accredited departments retain O
# for 창의입문설계 because it is mandatory within their accreditation programs.
FIRST_YEAR_OPTIONAL = {
    ("화학생명소재전공", "창의입문설계"),
}

GROUPS = [
    ("9개 전공 공통", COMMON_FILL, COMMON_DARK, COMMON),
    ("기초과학", SCIENCE_FILL, SCIENCE_DARK, SCIENCE),
    ("설계", DESIGN_FILL, DESIGN_DARK, DESIGN),
    ("MSC", MSC_FILL, MSC_DARK, MSC),
]


def set_cell_border(cell, *, color, size=20, edges=("top", "left", "bottom", "right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in edges:
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = base.RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(9)
    section.bottom_margin = Mm(9)
    section.left_margin = Mm(9)
    section.right_margin = Mm(9)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(4)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("DX대학 자율전공 교육과정 | 교과목 교집합표")
    base.set_font(run, size=8, bold=True, color=TEXT_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("근거: 2026학년도 교육과정 운영기준 및 편성표  |  2026.08.09  |  1 / 1")
    base.set_font(run, size=7.5, color=TEXT_GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("DX대학 9개 전공의 1학년 교과목 교집합")
    base.set_font(run, size=15.5, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("O는 1학년 필수, △는 1학년 선택, #-#는 기타 학년 필수(학년-학기)를 뜻합니다. 전공·계열 특화 과목은 교과목명을 직접 표시했습니다.")
    base.set_font(run, size=8.5, color=TEXT_GRAY)

    regular_courses = [course for _, _, _, courses in GROUPS for course in courses]
    total_cols = 1 + len(regular_courses) + 1
    table = doc.add_table(rows=2 + len(base.DEPARTMENTS), cols=total_cols)

    department_width = 2100
    regular_width = 550
    special_width = 2700
    widths = [department_width] + [regular_width] * len(regular_courses) + [special_width]
    base.geometry(table, widths)
    base.borders(table, color=GRID, size=5)

    dept_header = table.cell(0, 0).merge(table.cell(1, 0))
    base.shade(dept_header, NAVY)
    base.write(dept_header, "학과(전공)", size=8.2, bold=True, color=WHITE)

    col_index = 1
    course_meta = []
    for group_name, fill, color, courses in GROUPS:
        start = col_index
        end = col_index + len(courses) - 1
        group_cell = table.cell(0, start)
        if end > start:
            group_cell = group_cell.merge(table.cell(0, end))
        base.shade(group_cell, fill)
        base.write(group_cell, group_name, size=8.0, bold=True, color=color)
        for label, key in courses:
            cell = table.cell(1, col_index)
            base.shade(cell, fill)
            base.write(cell, label, size=6.8, bold=True, color=color)
            course_meta.append((key, fill, color))
            col_index += 1

    special_col = total_cols - 1
    special_group = table.cell(0, special_col)
    base.shade(special_group, SPECIAL_FILL)
    base.write(special_group, "전공·계열 특화", size=8.0, bold=True, color=SPECIAL_DARK)
    special_header = table.cell(1, special_col)
    base.shade(special_header, SPECIAL_FILL)
    base.write(special_header, "1학년 편성 교과목", size=7.4, bold=True, color=SPECIAL_DARK)

    accredited_rows = [
        row_index
        for row_index, (department, _) in enumerate(base.DEPARTMENTS, start=2)
        if "(인증)" in department
    ]

    for row_index, (department, subjects) in enumerate(base.DEPARTMENTS, start=2):
        dept_cell = table.cell(row_index, 0)
        base.shade(dept_cell, "F4F6F8" if row_index % 2 == 0 else WHITE)
        base.write(dept_cell, department, size=7.4, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT)
        if row_index in accredited_rows:
            # Draw one outer boundary around the union of all four accredited cells.
            # Internal row separators remain the table's normal thin grid lines.
            edges = ["left", "right"]
            if row_index == accredited_rows[0]:
                edges.append("top")
            if row_index == accredited_rows[-1]:
                edges.append("bottom")
            set_cell_border(dept_cell, color=ACCREDITED_BORDER, size=18, edges=edges)

        for offset, (key, fill, color) in enumerate(course_meta, start=1):
            cell = table.cell(row_index, offset)
            if (department, key) in FIRST_YEAR_OPTIONAL:
                base.shade(cell, fill)
                base.write(cell, "△", size=9.0, bold=True, color=color)
            elif key in subjects:
                base.shade(cell, fill)
                base.write(cell, "O", size=9.0, bold=True, color=color)
            elif key in LATER_REQUIRED.get(department, {}):
                base.shade(cell, fill)
                base.write(cell, LATER_REQUIRED[department][key], size=8.1, bold=True, color=color)
            else:
                base.shade(cell, "FAFBFC" if row_index % 2 == 0 else WHITE)
                base.write(cell, "", size=8)

        special_names = [key for key in SPECIAL_KEYS if key in subjects]
        special_cell = table.cell(row_index, special_col)
        base.shade(special_cell, SPECIAL_FILL if special_names else ("FAFBFC" if row_index % 2 == 0 else WHITE))
        base.write(
            special_cell,
            " · ".join(special_names),
            size=7.0,
            bold=bool(special_names),
            color=SPECIAL_DARK,
            align=WD_ALIGN_PARAGRAPH.LEFT if special_names else WD_ALIGN_PARAGRAPH.CENTER,
        )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("표시 기준  ")
    base.set_font(run, size=8, bold=True, color=BLUE)
    run = note.add_run("O는 1학년 필수, △는 1학년 선택, #-#는 기타학년 필수(학년-학기)입니다. 전공·계열 특화 칸은 1학년 편성 교과목명을 표시합니다. ")
    base.set_font(run, size=8, color=TEXT_GRAY)
    run = note.add_run("설계 글자색 외곽선")
    base.set_font(run, size=8, bold=True, color=ACCREDITED_BORDER)
    run = note.add_run("은 현재 확인된 공학인증 유지 4개 전공을 표시합니다.")
    base.set_font(run, size=8, color=TEXT_GRAY)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
