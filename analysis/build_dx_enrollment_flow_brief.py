"""Build a one-page (A4 PORTRAIT) enrollment-flow mind map for the 2027 restructuring.

Layout (four entry units on one shared top line):
  1) 첨단융합대학(AX)  2) 미래융합대학(DX)  3) 자율전공학부 유형 I  4) 기타 모집단위

Key relationships shown:
- Root "전체 신입생 입학" at the top center feeds every unit.
- 첨단융합대학 / 미래융합대학: admission SPLITS into two paths that both end in the
  college's own majors — (a) 자율전공 유형 II (1st-year detour, curved arrow at year end)
  and (b) 소속 학과·전공 직접입학 (straight, immediate).
- 자율전공학부 유형 I box mirrors the 유형 II cell (same size / cell position); its dashed
  arrows leave the BOTTOM of the box and disperse to the majors of the other three units.
- Every unit lists its member majors; the major-list tops are aligned across units.
- No head-count numbers on nodes; edge thickness only encodes relative intake scale.

College rosters follow the official 전공교육과정편성표(2026) 단과대학 편제(제4조).
"""

from __future__ import annotations

from math import atan2, cos, sin
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "DX대학_자율전공_1학년_교육과정_회의브리프_1p.docx"
ASSET_DIR = ROOT / "output" / "assets"
DIAGRAM = ASSET_DIR / "dx_enrollment_flow_mindmap.png"

NAVY = "17365D"
BLUE = "2E74B5"      # 미래융합대학 (DX)
TEAL = "237A78"      # 자율전공 유형 I
GOLD = "B07A12"      # 자율전공 유형 II
PURPLE = "6B4EA0"    # 첨단융합대학 (AX)
GRAYU = "5A6675"     # 기타 모집단위
INK = "1D2733"
GRAY = "647180"
LIGHT_BLUE = "EAF2FA"
LIGHT_TEAL = "E8F4F2"
LIGHT_GOLD = "FFF4D8"
LIGHT_PURPLE = "F0ECF9"
LIGHT_GRAY = "F1F3F6"
LINE = "CAD3DE"
WHITE = "FFFFFF"


def _first_existing(candidates):
    for path, index in candidates:
        if Path(path).exists():
            return path, index
    return candidates[0]


FONT_REG_PATH, FONT_REG_IDX = _first_existing([
    (r"C:\Windows\Fonts\malgun.ttf", 0),
    ("/usr/share/fonts/truetype/nanum/NanumGothic.ttf", 0),
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 1),
])
FONT_BOLD_PATH, FONT_BOLD_IDX = _first_existing([
    (r"C:\Windows\Fonts\malgunbd.ttf", 0),
    ("/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf", 0),
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc", 1),
])


def rgb(value: str) -> tuple[int, int, int]:
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def fnt(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if bold:
        return ImageFont.truetype(FONT_BOLD_PATH, size=size, index=FONT_BOLD_IDX)
    return ImageFont.truetype(FONT_REG_PATH, size=size, index=FONT_REG_IDX)


def rounded(draw, box, fill, outline=None, width=3, radius=24):
    draw.rounded_rectangle(
        box, radius=radius,
        fill=rgb(fill) if fill else None,
        outline=rgb(outline) if outline else None,
        width=width,
    )


def text_left(draw, xy, text, *, size, color=INK, bold=False, spacing=7):
    draw.multiline_text(xy, text, font=fnt(size, bold), fill=rgb(color), spacing=spacing)


def text_center(draw, box, text, *, size, color=INK, bold=False, spacing=6):
    x1, y1, x2, y2 = box
    font = fnt(size, bold)
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(
        ((x1 + x2 - w) / 2, (y1 + y2 - h) / 2 - bbox[1]),
        text, font=font, fill=rgb(color), spacing=spacing, align="center",
    )


def bezier(start, c1, c2, end, steps=90):
    pts = []
    for i in range(steps + 1):
        t = i / steps
        u = 1 - t
        x = u**3 * start[0] + 3 * u**2 * t * c1[0] + 3 * u * t**2 * c2[0] + t**3 * end[0]
        y = u**3 * start[1] + 3 * u**2 * t * c1[1] + 3 * u * t**2 * c2[1] + t**3 * end[1]
        pts.append((x, y))
    return pts


def arrowhead(draw, pts, color, width):
    x, y = pts[-1]
    px, py = pts[-4]
    ang = atan2(y - py, x - px)
    length = max(22, width * 2.0)
    spread = max(13, width * 1.0)
    bx, by = x - length * cos(ang), y - length * sin(ang)
    left = (bx + spread * cos(ang + 1.5708), by + spread * sin(ang + 1.5708))
    right = (bx + spread * cos(ang - 1.5708), by + spread * sin(ang - 1.5708))
    draw.polygon([(x, y), left, right], fill=rgb(color))


def curved_arrow(draw, start, c1, c2, end, *, color, width, dashed=False, arrow=True):
    pts = bezier(start, c1, c2, end)
    if dashed:
        for i in range(0, len(pts) - 1, 4):
            draw.line(pts[i : min(i + 3, len(pts))], fill=rgb(color), width=width, joint="curve")
    else:
        draw.line(pts, fill=rgb(color), width=width, joint="curve")
    if arrow:
        arrowhead(draw, pts, color, width)


def cell(draw, box, *, title, sub, accent, fill):
    """A rounded entry cell with a title line and a sub line (유형 II / 유형 I / 직접입학)."""
    rounded(draw, box, fill, outline=accent, width=3, radius=16)
    x1, y1, x2, y2 = box
    text_center(draw, (x1 + 8, y1 + 14, x2 - 8, y1 + 72), title, size=24, color=accent, bold=True)
    text_center(draw, (x1 + 8, y1 + 74, x2 - 8, y2 - 12), sub, size=18, color=GRAY, bold=True)


def major_card(draw, box, text, *, size=23, accent=LINE):
    rounded(draw, box, WHITE, outline=accent, width=2, radius=11)
    text_center(draw, box, text, size=size, color=INK, bold=True)


# ------------------------------- geometry ---------------------------------
W = 2400
MARGIN = 54
COL_W = 540
GUT = 44
C1 = MARGIN                    # 첨단융합대학 (AX)
C2 = C1 + COL_W + GUT          # 미래융합대학 (DX)
C3 = C2 + COL_W + GUT          # 자율전공학부 유형 I
C4 = C3 + COL_W + GUT          # 기타 모집단위

ROOT_BOX = (930, 96, 1470, 306)
TOP_Y = 520                    # unit box top (shared line)
CELL_W = 300
CELL_H = 130
T2_TOP = TOP_Y + 150           # 유형 II / 유형 I cell top (shared line)
DEPT_TOP = TOP_Y + 340         # major-list box top (shared line)
CARD_TOP = DEPT_TOP + 96       # first major card top (shared line, req 4)
CARD_STEP = 132
CARD_H = 86

AX_MAJORS = [
    "기계공학전공", "기계시스템공학전공", "스마트모빌리티전공",
    "반도체시스템전공", "전자시스템전공",
    "소프트웨어전공", "인공지능공학전공", "컴퓨터공학전공",
]
DX_MAJORS = [
    "건축공학전공", "토목공학전공", "산업공학전공", "수리빅데이터전공",
    "고분자공학전공", "신소재공학전공", "소재디자인공학전공",
    "화학공학전공", "화학생명소재전공",
]
ETC_UNITS = [
    "에디슨칼리지(첨단산업융합학부)", "바이오메디컬공학과", "AX융합학과",
    "반도체소부장전공", "건축학과(5년제)", "광시스템공학과",
    "환경공학과", "IT융합학과", "경영학과", "글로컬융합학부(REC)",
]


def col_cx(x_left):
    return x_left + COL_W / 2


def dept_box_bottom(n):
    return CARD_TOP + n * CARD_STEP + 14


def draw_major_list(draw, x_left, majors, *, accent, dept_title, size=23):
    box = (x_left + 32, DEPT_TOP, x_left + COL_W - 32, dept_box_bottom(len(majors)))
    rounded(draw, box, WHITE, outline=accent, width=3, radius=20)
    text_center(draw, (box[0] + 14, DEPT_TOP + 16, box[2] - 14, DEPT_TOP + 80), dept_title, size=23, color=accent, bold=True)
    gy = CARD_TOP
    for name in majors:
        major_card(draw, (box[0] + 22, gy, box[2] - 22, gy + CARD_H), name, size=size, accent=accent)
        gy += CARD_STEP
    return box


def build_diagram():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    etc_bottom = dept_box_bottom(len(ETC_UNITS))
    H = int(etc_bottom + 150)
    img = Image.new("RGB", (W, H), rgb(WHITE))
    draw = ImageDraw.Draw(img)

    # ----- root node -----
    rcx = (ROOT_BOX[0] + ROOT_BOX[2]) / 2
    r_bottom = ROOT_BOX[3]

    # ----- admission edges from root to each unit (thickness = relative scale) -----
    w_ax, w_dx, w_t1, w_etc = 26, 24, 13, 16
    ax_cx, dx_cx, t1_cx, etc_cx = col_cx(C1), col_cx(C2), col_cx(C3), col_cx(C4)
    for cx, wdt, col in [(ax_cx, w_ax, PURPLE), (dx_cx, w_dx, BLUE), (t1_cx, w_t1, TEAL), (etc_cx, w_etc, GRAYU)]:
        curved_arrow(draw, (rcx, r_bottom - 6), (rcx + (cx - rcx) * 0.25, r_bottom + 70),
                     (cx, TOP_Y - 90), (cx, TOP_Y - 4), color=col, width=wdt)

    rounded(draw, ROOT_BOX, NAVY, radius=48)
    text_center(draw, (ROOT_BOX[0], ROOT_BOX[1] + 26, ROOT_BOX[2], ROOT_BOX[1] + 104), "전체 신입생", size=32, color=WHITE, bold=True)
    text_center(draw, (ROOT_BOX[0], ROOT_BOX[1] + 104, ROOT_BOX[2], ROOT_BOX[3] - 16), "입학", size=46, color=WHITE, bold=True)

    # ============================ COLLEGE COLUMNS (AX, DX) ============================
    def college(x_left, *, title, sub, accent, fill, majors, dept_title):
        box = (x_left, TOP_Y, x_left + COL_W, dept_box_bottom(len(majors)) + 34)
        rounded(draw, box, fill, outline=accent, width=5, radius=28)
        text_center(draw, (x_left + 18, TOP_Y + 18, x_left + COL_W - 18, TOP_Y + 82), title, size=30, color=accent, bold=True)
        text_center(draw, (x_left + 18, TOP_Y + 84, x_left + COL_W - 18, TOP_Y + 120), sub, size=18, color=GRAY, bold=True)

        # two entry cells: 유형 II (left) + 직접입학 (right)
        t2 = (x_left + 40, T2_TOP, x_left + 40 + CELL_W, T2_TOP + CELL_H)
        cell(draw, t2, title="자율전공 유형 II", sub="1학년 과정만 운영", accent=GOLD, fill=LIGHT_GOLD)
        dr = (x_left + 40 + CELL_W + 18, T2_TOP + 6, x_left + COL_W - 40, T2_TOP + CELL_H - 6)
        cell(draw, dr, title="직접입학", sub="소속 학과·전공", accent=accent, fill=WHITE)

        deptbox = draw_major_list(draw, x_left, majors, accent=accent, dept_title=dept_title)

        # 유형 II students pick a major at YEAR END -> curved (left lane)
        curved_arrow(draw, ((t2[0] + t2[2]) / 2, t2[3]),
                     (t2[0] + 20, t2[3] + 46), (deptbox[0] + 60, DEPT_TOP - 40),
                     (deptbox[0] + 96, DEPT_TOP + 2), color=GOLD, width=8)
        # direct-admission flows straight into the majors (right lane)
        curved_arrow(draw, ((dr[0] + dr[2]) / 2, dr[3]),
                     ((dr[0] + dr[2]) / 2, dr[3] + 30), (deptbox[2] - 96, DEPT_TOP - 40),
                     (deptbox[2] - 96, DEPT_TOP + 2), color=accent, width=10)
        return box, deptbox

    ax_box, ax_dept = college(C1, title="첨단융합대학", sub="(AX대학)", accent=PURPLE, fill=LIGHT_PURPLE,
                              majors=AX_MAJORS, dept_title="첨단융합대학 소속 전공")
    dx_box, dx_dept = college(C2, title="미래융합대학", sub="(DX대학)", accent=BLUE, fill=LIGHT_BLUE,
                              majors=DX_MAJORS, dept_title="미래융합대학 소속 전공")

    # ============================ 자율전공학부 유형 I (col 3) ============================
    t1_box = (C3, TOP_Y, C3 + COL_W, DEPT_TOP + 120)
    rounded(draw, t1_box, LIGHT_TEAL, outline=TEAL, width=5, radius=28)
    text_center(draw, (C3 + 18, TOP_Y + 30, C3 + COL_W - 18, TOP_Y + 104), "자율전공학부", size=30, color=TEAL, bold=True)
    text_center(draw, (C3 + 18, TOP_Y + 104, C3 + COL_W - 18, TOP_Y + 146), "유형 I", size=24, color=TEAL, bold=True)
    # 유형 I cell — SAME size and cell position as the 유형 II cell (req 2)
    t1_cell = (C3 + 40, T2_TOP, C3 + 40 + CELL_W, T2_TOP + CELL_H)
    cell(draw, t1_cell, title="자율전공 유형 I", sub="1학년 과정만 운영", accent=TEAL, fill=LIGHT_TEAL)
    text_center(draw, (C3 + 34, T2_TOP + CELL_H + 20, C3 + COL_W - 34, DEPT_TOP + 96),
                "1학년 말\n대학 전체 학과·전공에서\n선택", size=21, color=INK, bold=True)

    # ============================ 기타 모집단위 (col 4) ============================
    etc_box = (C4, TOP_Y, C4 + COL_W, dept_box_bottom(len(ETC_UNITS)) + 34)
    rounded(draw, etc_box, LIGHT_GRAY, outline=GRAYU, width=5, radius=28)
    text_center(draw, (C4 + 18, TOP_Y + 18, C4 + COL_W - 18, TOP_Y + 82), "기타 모집단위", size=30, color=GRAYU, bold=True)
    text_center(draw, (C4 + 18, TOP_Y + 84, C4 + COL_W - 18, TOP_Y + 120), "(그 외 단과대학·학과)", size=18, color=GRAY, bold=True)
    etc_dr = (C4 + 90, T2_TOP + 20, C4 + COL_W - 90, T2_TOP + CELL_H - 20)
    cell(draw, etc_dr, title="직접입학", sub="소속 학과·전공", accent=GRAYU, fill=WHITE)
    etc_dept = draw_major_list(draw, C4, ETC_UNITS, accent=GRAYU, dept_title="기타 모집단위(학과·학부)", size=18)
    curved_arrow(draw, ((etc_dr[0] + etc_dr[2]) / 2, etc_dr[3]),
                 ((etc_dr[0] + etc_dr[2]) / 2, etc_dr[3] + 30), (etc_cx, DEPT_TOP - 40),
                 (etc_cx, DEPT_TOP + 2), color=GRAYU, width=10)

    # ============================ 유형 I dashed dispersal (from BOTTOM, req 3) ============
    # All three leave the BOTTOM of the 유형 I box, rise to a clear horizontal channel that
    # sits ABOVE every aligned dept-title, sweep sideways, then dive into a target major list.
    b_y = t1_box[3]
    # -> 미래융합대학 (adjacent left)
    curved_arrow(draw, (C3 + 60, b_y), (C3 - 10, DEPT_TOP - 40), (dx_dept[2] - 40, DEPT_TOP - 34),
                 (dx_dept[2] - 66, DEPT_TOP + 2), color=TEAL, width=6, dashed=True)
    # -> 기타 모집단위 (adjacent right)
    curved_arrow(draw, (C3 + COL_W - 60, b_y), (C3 + COL_W + 10, DEPT_TOP - 40), (etc_dept[0] + 40, DEPT_TOP - 34),
                 (etc_dept[0] + 66, DEPT_TOP + 2), color=TEAL, width=6, dashed=True)
    # -> 첨단융합대학 (far left; horizontal run held at DEPT_TOP-56, clear of all dept titles,
    #    diving into the CENTRE-top of the 첨단융합 major list so the target is unambiguous)
    ax_cx2 = (ax_dept[0] + ax_dept[2]) / 2
    curved_arrow(draw, (C3 + 40, b_y), (C2 + 60, DEPT_TOP - 56), (ax_cx2 + 40, DEPT_TOP - 56),
                 (ax_cx2, DEPT_TOP - 1), color=TEAL, width=6, dashed=True)

    # ============================ legend ============================
    leg = (MARGIN, H - 116, W - MARGIN, H - 44)
    rounded(draw, leg, "FFF8E8", outline=GOLD, width=3, radius=16)
    text_left(draw, (leg[0] + 30, leg[1] + 14), "읽는 법", size=21, color=GOLD, bold=True)
    text_left(draw, (leg[0] + 160, leg[1] + 14), "입학선 두께 = 상대 규모   ·   굵은 실선·직접입학 → 소속 학과·전공으로 곧바로 입학", size=18, color=INK, bold=True)
    text_left(draw, (leg[0] + 160, leg[1] + 44), "금색 곡선 = 유형 II 학생의 1학년 말 소속 전공 선택   ·   청록 점선 = 유형 I 학생의 1학년 말 전공 분산", size=18, color=GRAY)

    img.save(DIAGRAM, quality=95)
    return W, H


def set_run_font(run, *, size, bold=False, color=INK):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_bottom_border(paragraph, color=BLUE, size=12):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def build_docx():
    w_px, h_px = build_diagram()
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.18)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("2027 학사구조 개편 안내  |  신입생 소속·전공 선택 흐름\n"), size=8.5, bold=True, color=GRAY)
    set_run_font(header.add_run("github.com/cfms-lab/KIT_DX_2027"), size=7.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.add_run("국립금오공과대학교  ·  전공교육과정편성표(2026) 단과대학 편제 기준  |  2026.08.09  |  1 / 1"), size=8, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(0)
    set_run_font(kicker.add_run("공통 설명자료  ·  입학 단위에서 1학년 말 전공 선택까지"), size=9.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    set_run_font(title.add_run("전체 신입생은 어떤 경로로 단과대학과 전공에 진입하는가"), size=19, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(5)
    set_run_font(subtitle.add_run("첨단융합대학·미래융합대학·자율전공 유형 I·기타 모집단위의 입학 경로와 1학년 말 전공 선택"), size=10, bold=True, color=GRAY)
    add_bottom_border(subtitle)

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_before = Pt(3)
    pic.paragraph_format.space_after = Pt(2)
    pic.add_run().add_picture(str(DIAGRAM), width=Inches(7.2))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(1)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    set_run_font(note.add_run("해석  "), size=8.3, bold=True, color=BLUE)
    set_run_font(note.add_run("첨단융합대학·미래융합대학은 ‘자율전공 유형 II’와 ‘소속 학과·전공 직접입학’ 두 경로로 입학하며, 두 경로 모두 해당 대학의 전공으로 귀결됨.  "), size=8.3, color=GRAY)
    set_run_font(note.add_run("유형 I  "), size=8.3, bold=True, color=TEAL)
    set_run_font(note.add_run("자율전공 유형 I 학생은 1학년 말 대학 전체 전공에서 선택하며(청록 점선), 정원 숫자는 도식에 표기하지 않고 선 두께로 상대 규모만 나타냄."), size=8.3, color=GRAY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"{OUT}  (canvas {w_px}x{h_px})")


if __name__ == "__main__":
    build_docx()
