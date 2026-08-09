"""Build a one-page decision brief for the DX autonomous-major curriculum meeting."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "DX대학_자율전공_1학년_교육과정_회의브리프_1p.docx"

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "DCEAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9DEE5"
TEXT_GRAY = "4A5560"
PALE_YELLOW = "FFF4CC"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360
TABLE_TEXT_SIZE = 9.2  # Named compact-table override for the one-page form factor.


def set_repeat_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def write_cell(cell, text, *, bold=False, color="000000", size=TABLE_TEXT_SIZE, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_repeat_font(r, size=size, bold=bold, color=color)


def add_section_label(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_repeat_font(r, size=11.5, bold=True, color=NAVY)
    return p


def add_compact_note(doc, lead, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(lead)
    set_repeat_font(r, size=9.2, bold=True, color=BLUE)
    r = p.add_run(text)
    set_repeat_font(r, size=9.2, color=TEXT_GRAY)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.40)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.16)
    section.footer_distance = Inches(0.16)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    # Running header/footer: restrained memo furniture, no decorative rule.
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("DX대학 자율전공 교육과정 | 회의 브리프")
    set_repeat_font(hr, size=8.5, bold=True, color=TEXT_GRAY)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("근거: 2026학년도 교육과정 운영기준 및 편성표  |  2026.08.09  |  1 / 1")
    set_repeat_font(fr, size=8, color=TEXT_GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1)
    kr = kicker.add_run("월요일 학과장 회의 · 사실 확인용")
    set_repeat_font(kr, size=9, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run("DX대학 자율전공 1학년 교육과정")
    set_repeat_font(tr, size=17.5, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(5)
    sr = sub.add_run("9개 전공의 현재 과목 분포와 결정이 필요한 범위")
    set_repeat_font(sr, size=11, bold=True, color=TEXT_GRAY)

    callout = doc.add_table(rows=1, cols=2)
    set_table_geometry(callout, [1200, 8160])
    set_table_borders(callout, color="E5C65E", size=7)
    set_cell_shading(callout.cell(0, 0), PALE_YELLOW)
    set_cell_shading(callout.cell(0, 1), PALE_YELLOW)
    write_cell(callout.cell(0, 0), "자료 성격", bold=True, color=NAVY, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    write_cell(
        callout.cell(0, 1),
        "특정 논점·편성안을 미리 정하지 않은 과목 분포 및 미이수 시 보완 가능성 자료",
        bold=True,
        size=9.5,
    )

    add_section_label(doc, "1. 전공군별 현재 1학년 구조")
    clusters = doc.add_table(rows=5, cols=3)
    set_table_geometry(clusters, [2200, 2200, 4960])
    set_table_borders(clusters)
    headers = ["전공군", "고정학점", "두드러지는 1학년 과목"]
    for i, text in enumerate(headers):
        set_cell_shading(clusters.cell(0, i), PALE_BLUE)
        write_cell(clusters.cell(0, i), text, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    cluster_rows = [
        ("건축·토목", "각 30", "일반물리1·2, 일반화학1 + 건축토목환경공학개론"),
        ("산업·수리빅데이터", "26 / 23", "확률및통계, 컴퓨터프로그래밍언어, 빅데이터의세계"),
        ("인증 소재·화학 4개", "각 33", "일반물리·화학 1·2와 실험 + 창의입문설계; 공학인증 유지"),
        ("화학생명소재", "33", "동일한 물리·화학·실험 및 창의입문설계 편성; 공학인증은 아님"),
    ]
    for r_idx, row in enumerate(cluster_rows, 1):
        for c_idx, text in enumerate(row):
            write_cell(clusters.cell(r_idx, c_idx), text, bold=(c_idx == 0), color=NAVY if c_idx == 0 else "000000")

    add_section_label(doc, "2. 과목군별 적용 범위")
    coverage = doc.add_table(rows=7, cols=4)
    set_table_geometry(coverage, [2380, 1180, 3000, 2800])
    set_table_borders(coverage)
    coverage_headers = ["과목군", "전공 수", "현재 포함 전공", "확인해야 할 점"]
    for i, text in enumerate(coverage_headers):
        set_cell_shading(coverage.cell(0, i), PALE_BLUE)
        write_cell(coverage.cell(0, i), text, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    coverage_rows = [
        ("공통 5과목(14학점)", "9/9", "전체", "학기 배치는 서로 다름"),
        ("일반물리학1·2", "7/9", "건축·토목·소재/화학 5개", "산업·수리빅데이터에는 없음"),
        ("일반화학1", "7/9", "건축·토목·소재/화학 5개", "산업·수리빅데이터에는 없음"),
        ("일반화학2 + 물리·화학실험", "5/9", "소재/화학 5개", "건축·토목도 현재는 미편성"),
        ("창의입문설계", "5/9", "소재/화학 5개", "그중 4개는 공학인증 유지"),
        ("프로그래밍 / 확률통계", "3/9 · 2/9", "토목·산업·수리 / 산업·수리", "다른 전공에는 1학년 필수 아님"),
    ]
    for r_idx, row in enumerate(coverage_rows, 1):
        for c_idx, text in enumerate(row):
            write_cell(coverage.cell(r_idx, c_idx), text, bold=(c_idx == 0), color=NAVY if c_idx == 0 else "000000", size=8.8)

    add_section_label(doc, "3. 어느 과목을 필수화해도 함께 발생하는 영향")
    impacts = doc.add_table(rows=4, cols=3)
    set_table_geometry(impacts, [2300, 3530, 3530])
    set_table_borders(impacts)
    impact_headers = ["검토 과목군", "필수에 포함할 때", "필수에서 제외할 때"]
    for i, text in enumerate(impact_headers):
        set_cell_shading(impacts.cell(0, i), PALE_BLUE)
        write_cell(impacts.cell(0, i), text, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    impact_rows = [
        ("일반물리·일반화학", "산업·수리빅데이터 학생에게 현재 없는 과학 과목 부담", "다수 공학 전공 진입 시 선수지식·추가이수 검토 필요"),
        ("물리·화학실험", "건축·토목·산업·수리에는 현재 없는 실험학점 부담", "소재·화학계 진입 시 실험 교과 연계성 확인 필요"),
        ("창의입문설계", "비인증 전공 학생에게 3학점 공통 부담", "공학인증 4개 전공 진입자의 설계학점 보완 가능성"),
    ]
    for r_idx, row in enumerate(impact_rows, 1):
        for c_idx, text in enumerate(row):
            write_cell(impacts.cell(r_idx, c_idx), text, bold=(c_idx == 0), color=NAVY if c_idx == 0 else "000000", size=8.9)

    add_section_label(doc, "4. 회의에서 정해야 할 범위")
    decision_table = doc.add_table(rows=2, cols=2)
    set_table_geometry(decision_table, [4680, 4680])
    set_table_borders(decision_table)
    decision_items = [
        "1  필수 선정 기준: 교집합·진입 호환성·기초소양·전공탐색의 우선순위",
        "2  물리·화학·실험·설계·프로그래밍/통계의 필수·선택·분기 여부",
        "3  1학년 총 필수학점 상한과 전공 진입 후 허용할 추가이수 범위",
        "4  EA0001·HA0001·KA0001 및 공통 기초과목의 동일인정 기준",
    ]
    for idx, item in enumerate(decision_items):
        row, col = divmod(idx, 2)
        if row == 0:
            set_cell_shading(decision_table.cell(row, col), LIGHT_GRAY)
        write_cell(decision_table.cell(row, col), item, bold=True, color=NAVY, size=8.7)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
